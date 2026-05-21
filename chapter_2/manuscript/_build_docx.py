"""Convert chapter_2.md to Manning-styled chapter_2.docx.

Mapping (per agents/research_and_draft_chapter/SKILL.md §10):
  # <N> ... + # <Title>           -> Title = "<N>", Subtitle = "<Title>"
  ## <N>.<m> ...                  -> Heading 1
  ### <N>.<m>.<k> ...             -> Heading 2
  #### ...                        -> Heading 3
  **Listing X.Y** caption.        -> Heading 4 (caption line)
  **Figure X.Y** caption.         -> Heading 6 (caption line)
  **Table X.Y** caption.          -> Heading 5 (caption line)
  ```python ... ```               -> Normal para, Courier New 9pt, light grey shading
  > **LABEL — title.** body       -> Normal para with bold lead and parchment shading
  - item                          -> Normal para indented, '•' lead
  | ... | ... |                   -> docx table with manual single-line 888888 borders
  ![Figure X.Y caption](path)     -> Heading 6 caption + embedded image (if exists)

Run: python _build_docx.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

MANUSCRIPT = Path(__file__).parent / "chapter_2.md"
TEMPLATE = Path(__file__).parent.parent.parent / "chapter_template" / "chapter_template.docx"
FIGURES_DIR = Path(__file__).parent.parent / "figures" / "diagrams"
OUTPUT = Path(__file__).parent / "chapter_2.docx"

CODE_FONT = "Courier New"
CODE_SIZE = Pt(9)
CODE_SHADING = "EEEEEE"
CALLOUT_SHADING = "F5F0E1"  # parchment
BORDER_COLOR = "888888"


def add_shading(paragraph, color_hex: str) -> None:
    """Apply a background shading color to a paragraph."""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    p_pr.append(shd)


def add_table_borders(table, color_hex: str = BORDER_COLOR) -> None:
    """Manually add single-line borders to every cell in a table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), color_hex)
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)


def parse_inline(paragraph, text: str) -> None:
    """Render markdown inline formatting (bold, italic, code) into runs."""
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|__[^_]+__|"      # bold
        r"\*[^*]+\*|_[^_]+_|"             # italic
        r"`[^`]+`)",                      # inline code
    )
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**") or token.startswith("__"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = CODE_FONT
            run.font.size = Pt(10)
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_code_block(doc, code_lines: list[str]) -> None:
    """Render a fenced code block as a shaded monospace paragraph."""
    para = doc.add_paragraph()
    add_shading(para, CODE_SHADING)
    for i, line in enumerate(code_lines):
        if i > 0:
            run = para.add_run()
            run.add_break()
        run = para.add_run(line)
        run.font.name = CODE_FONT
        run.font.size = CODE_SIZE


def add_callout_block(doc, lines: list[str]) -> None:
    """Render a blockquote callout with parchment shading and bold lead."""
    text = " ".join(line.lstrip("> ").strip() for line in lines)
    para = doc.add_paragraph()
    add_shading(para, CALLOUT_SHADING)
    parse_inline(para, text)


def add_bullet(doc, text: str) -> None:
    """Render a bullet line as a plain paragraph with '•' lead."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    run = para.add_run("• ")
    parse_inline(para, text)


def add_table(doc, rows: list[list[str]]) -> None:
    """Render a markdown table as a docx table with manual borders."""
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            parse_inline(para, cell_text)
            if r == 0:
                for run in para.runs:
                    run.bold = True
    add_table_borders(table)


def parse_table_lines(lines: list[str]) -> list[list[str]]:
    """Convert raw markdown table lines into a 2D list of cell strings."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if all(re.fullmatch(r":?-+:?", c) for c in cells):
            continue  # alignment row
        rows.append(cells)
    return rows


def add_figure(doc, caption: str, image_path: str) -> None:
    """Add a Heading 6 caption and embed the image if it exists."""
    doc.add_heading(caption, level=6)
    img = (FIGURES_DIR / Path(image_path).name)
    if img.exists():
        doc.add_picture(str(img), width=Inches(5.5))
    else:
        placeholder = doc.add_paragraph(
            f"[Figure placeholder: {image_path} not yet rendered.]"
        )
        for run in placeholder.runs:
            run.italic = True


def convert(markdown_path: Path, output_path: Path, template_path: Path) -> None:
    if template_path.exists():
        doc = Document(str(template_path))
        # Clear template body content
        for elem in list(doc.element.body):
            doc.element.body.remove(elem)
        doc.element.body.append(OxmlElement("w:sectPr"))
    else:
        doc = Document()

    with markdown_path.open() as f:
        lines = f.read().splitlines()

    title_set = subtitle_set = False
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Chapter number (first H1 = title; second H1 = subtitle)
        if stripped.startswith("# ") and not stripped.startswith("## "):
            heading_text = stripped[2:].strip()
            if not title_set:
                doc.add_heading(heading_text, level=0)  # Title style
                title_set = True
            elif not subtitle_set:
                # Subtitle: render as bold large text below title
                para = doc.add_paragraph()
                run = para.add_run(heading_text)
                run.bold = True
                run.font.size = Pt(24)
                subtitle_set = True
            else:
                doc.add_heading(heading_text, level=1)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
            i += 1
            continue

        # Image / figure
        img_match = re.match(r"!\[(.+?)\]\((.+?)\)", stripped)
        if img_match:
            add_figure(doc, img_match.group(1), img_match.group(2))
            i += 1
            continue

        # Fenced code block
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1  # skip the closing ```
            continue

        # Blockquote callout (one or more consecutive lines)
        if stripped.startswith(">"):
            callout_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                callout_lines.append(lines[i])
                i += 1
            add_callout_block(doc, callout_lines)
            continue

        # Listing / Figure / Table caption (bold paragraph that starts with **Listing/**Figure/**Table)
        cap = re.match(r"\*\*(Listing|Figure|Table) (\d+\.\d+)\s*([^*]+?)\.?\*\*\s*$",
                       stripped)
        if cap:
            kind = cap.group(1)
            num = cap.group(2)
            title = cap.group(3).strip()
            level = {"Listing": 4, "Table": 5, "Figure": 6}[kind]
            doc.add_heading(f"{kind} {num} {title}", level=level)
            i += 1
            continue

        # Table
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table_lines(table_lines)
            add_table(doc, rows)
            continue

        # Bullet list
        if stripped.startswith("- "):
            add_bullet(doc, stripped[2:].strip())
            i += 1
            continue

        # Numbered list
        num_match = re.match(r"(\d+)\.\s+(.+)", stripped)
        if num_match:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            para.add_run(f"{num_match.group(1)}. ")
            parse_inline(para, num_match.group(2))
            i += 1
            continue

        # Annotation legend line: "#A description"
        ann_match = re.match(r"#([A-Z])\s+(.+)", stripped)
        if ann_match:
            para = doc.add_paragraph()
            run = para.add_run(f"#{ann_match.group(1)} ")
            run.font.name = CODE_FONT
            run.font.bold = True
            parse_inline(para, ann_match.group(2))
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            i += 1
            continue

        # Plain paragraph (collect contiguous non-empty lines)
        para_lines = [stripped]
        i += 1
        while (
            i < len(lines)
            and lines[i].strip()
            and not lines[i].strip().startswith(("#", ">", "-", "|", "```", "!"))
            and not re.match(r"\*\*(Listing|Figure|Table) \d+\.\d+", lines[i].strip())
            and not re.match(r"#[A-Z]\s+", lines[i].strip())
            and not re.match(r"\d+\.\s+", lines[i].strip())
        ):
            para_lines.append(lines[i].strip())
            i += 1
        para = doc.add_paragraph()
        parse_inline(para, " ".join(para_lines))

    doc.save(str(output_path))


def inspect(output_path: Path) -> None:
    """Sanity-check the generated DOCX: counts of headings, images, tables."""
    doc = Document(str(output_path))
    counts: dict[str, int] = {}
    for p in doc.paragraphs:
        style = p.style.name
        counts[style] = counts.get(style, 0) + 1
    images = 0
    for shape in doc.inline_shapes:
        images += 1
    tables = len(doc.tables)
    print("=== DOCX inspection ===")
    print(f"Output: {output_path}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    for style, n in sorted(counts.items(), key=lambda kv: -kv[1]):
        print(f"  {style}: {n}")
    print(f"Inline images: {images}")
    print(f"Tables: {tables}")


if __name__ == "__main__":
    convert(MANUSCRIPT, OUTPUT, TEMPLATE)
    inspect(OUTPUT)
